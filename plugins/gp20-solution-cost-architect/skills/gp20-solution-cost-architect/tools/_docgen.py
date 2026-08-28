"""Shared Word rendering for the sample tenders. Dev tooling, not runtime."""

from __future__ import annotations

import sys as _sys
for _s in (_sys.stdout, _sys.stderr):
    try:
        if _s and getattr(_s, "encoding", "").lower().replace("-", "") != "utf8":
            _s.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x35, 0x56)
GREY = RGBColor(0x6B, 0x77, 0x85)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def para(doc, text="", *, size=10.5, bold=False, italic=False, color=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return p


def heading(doc, text, size=15):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p


def bullets(doc, items):
    for t in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(t).font.size = Pt(10.5)


def table(doc, rows, widths_cm):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].width = Cm(widths_cm[ci])
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            if ri == 0 or (ri > 0 and ci > 0):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.bold = (ri == 0) or (ci == 0 and ri > 0)
            if ri == 0:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade(cells[ci], "1F3556")
    return t


def render(spec: dict, out_path: str) -> str:
    """spec: {ref, title, control[], sections[], annex{}}"""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.1)
        sec.left_margin = sec.right_margin = Cm(2.1)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    para(doc, "AURELIAN GLOBAL HOLDINGS PLC", size=19, bold=True, color=NAVY, after=2)
    para(doc, "Group Procurement — Technology Services", size=10.5, color=GREY, after=16)
    para(doc, "REQUEST FOR PROPOSAL", size=17, bold=True, after=4)
    para(doc, spec["title"], size=13, after=14)
    table(doc, [["Field", "Detail"]] + spec["control"], [5.4, 10.6])
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    for sec in spec["sections"]:
        heading(doc, sec["heading"])
        for block in sec["body"]:
            kind = block[0]
            if kind == "p":
                para(doc, block[1], italic=len(block) > 2 and block[2] == "i")
            elif kind == "b":
                bullets(doc, block[1])
            elif kind == "t":
                table(doc, block[1], block[2])

    annex = spec["annex"]
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    heading(doc, annex["heading"])
    para(doc, annex["intro"])
    table(doc, annex["rows"], annex["widths"])
    para(doc, "")
    for note in annex["notes"]:
        para(doc, note, size=9.5, italic=note.startswith("Note"))

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
